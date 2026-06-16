from __future__ import annotations
from pathlib import Path
from uuid import uuid4

from novel_runtime.services.project_service import ProjectService


class ExportService:
    def __init__(self, project_service: ProjectService):
        self.project_service = project_service
        self._export_dir = Path(project_service.storage_base) / "exports"

    def _collect_chapters(self, project_id: str, chapter_range: list[int] | None) -> list[dict]:
        chapters = self.project_service.list_chapters(project_id)
        if chapter_range:
            chapters = [c for c in chapters if c.chapter_number in range(chapter_range[0], chapter_range[1] + 1)]
        result = []
        for c in chapters:
            if c.status not in ("approved", "locked"):
                continue
            project_path = self.project_service.get_project_path(project_id)
            final_path = project_path / "chapters" / f"chapter_{c.chapter_number:03d}" / "final.md"
            if not final_path.exists():
                continue
            text = final_path.read_text(encoding="utf-8")
            result.append({"chapter_number": c.chapter_number, "title": c.title or f"第{c.chapter_number}章", "text": text})
        return result

    def export_txt(self, project_id: str, chapter_range: list[int] | None = None, include_title: bool = True) -> Path:
        chapters = self._collect_chapters(project_id, chapter_range)
        lines = []
        for c in chapters:
            if include_title:
                lines.append(f"# {c['title']}\n\n")
            lines.append(c["text"])
            lines.append("\n\n---\n\n")
        task_id = f"export_{uuid4().hex[:12]}"
        self._export_dir.mkdir(parents=True, exist_ok=True)
        path = self._export_dir / f"{task_id}.txt"
        path.write_text("".join(lines), encoding="utf-8")
        return path

    def export_md(self, project_id: str, chapter_range: list[int] | None = None, include_title: bool = True) -> Path:
        chapters = self._collect_chapters(project_id, chapter_range)
        lines = []
        for i, c in enumerate(chapters):
            if include_title:
                lines.append(f"# {c['title']}\n\n")
            lines.append(c["text"])
            if i < len(chapters) - 1:
                lines.append("\n\n---\n\n")
        task_id = f"export_{uuid4().hex[:12]}"
        self._export_dir.mkdir(parents=True, exist_ok=True)
        path = self._export_dir / f"{task_id}.md"
        path.write_text("".join(lines), encoding="utf-8")
        return path

    def export_epub(self, project_id: str, chapter_range: list[int] | None = None, include_title: bool = True) -> Path:
        chapters = self._collect_chapters(project_id, chapter_range)
        task_id = f"export_{uuid4().hex[:12]}"
        self._export_dir.mkdir(parents=True, exist_ok=True)
        path = self._export_dir / f"{task_id}.epub"

        from ebooklib import epub
        project = self.project_service.get_project(project_id)
        book = epub.EpubBook()
        book.set_identifier(project_id)
        book.set_title(project.project_name)
        book.set_language("zh-CN")

        for c in chapters:
            title = c["title"] if include_title else ""
            content = f"<h1>{title}</h1>\n" + c["text"].replace("\n", "<br/>\n") if title else c["text"].replace("\n", "<br/>\n")
            chap = epub.EpubHtml(title=c["title"], file_name=f"chap_{c['chapter_number']:03d}.xhtml", lang="zh-CN")
            chap.content = f"<html><body>{content}</body></html>"
            book.add_item(chap)
            book.toc.append(epub.Link(f"chap_{c['chapter_number']:03d}.xhtml", c["title"], f"chap_{c['chapter_number']:03d}"))
            book.add_item(chap)
            book.spine.append(chap)

        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        epub.write_epub(str(path), book)
        return path

    def export_docx(self, project_id: str, chapter_range: list[int] | None = None, include_title: bool = True) -> Path:
        from docx import Document
        chapters = self._collect_chapters(project_id, chapter_range)
        task_id = f"export_{uuid4().hex[:12]}"
        self._export_dir.mkdir(parents=True, exist_ok=True)
        path = self._export_dir / f"{task_id}.docx"

        doc = Document()
        for c in chapters:
            if include_title:
                doc.add_heading(c["title"], level=1)
            for para in c["text"].split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()
        doc.save(str(path))
        return path
